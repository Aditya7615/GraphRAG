"""Document loaders: PDF, TXT, MD, DOCX -> `SourceDocument` with page fidelity."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from backend.ingestion.chunker import Page, SourceDocument, content_hash, stable_id

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".pdf", ".txt", ".md", ".markdown", ".docx"}

# Ligatures and hyphenation artefacts that PDF extraction loves to produce.
_LIGATURES = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl", "\u00a0": " "}


def clean_text(text: str) -> str:
    for bad, good in _LIGATURES.items():
        text = text.replace(bad, good)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)  # de-hyphenate across line breaks
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _load_pdf(path: Path) -> list[Page]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[Page] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except (OSError, ValueError) as exc:  # a single corrupt page must not kill the file
            logger.warning("Failed to extract page %s of %s: %s", i, path.name, exc)
            raw = ""
        text = clean_text(raw)
        if text:
            pages.append(Page(number=i, text=text))
    return pages


def _load_docx(path: Path) -> list[Page]:
    import docx

    document = docx.Document(str(path))
    blocks = [p.text for p in document.paragraphs if p.text.strip()]
    text = clean_text("\n\n".join(blocks))
    return [Page(number=1, text=text)] if text else []


def _load_plaintext(path: Path) -> list[Page]:
    text = clean_text(path.read_text(encoding="utf-8", errors="replace"))
    if not text:
        return []
    # Treat form feeds or "--- Page N ---" markers as page breaks when present,
    # otherwise the whole file is one logical page.
    parts = re.split(r"\f|\n-{3,}\s*Page\s+\d+\s*-{3,}\n", text)
    parts = [p.strip() for p in parts if p.strip()]
    return [Page(number=i, text=p) for i, p in enumerate(parts, start=1)]


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_plaintext,
    ".md": _load_plaintext,
    ".markdown": _load_plaintext,
}


def load_document(path: str | Path, doc_type: str | None = None) -> SourceDocument:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"No such document: {path}")
    suffix = path.suffix.lower()
    if suffix not in _LOADERS:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {sorted(SUPPORTED_SUFFIXES)}")

    pages = _LOADERS[suffix](path)
    if not pages:
        raise ValueError(f"No extractable text found in {path.name}")

    body = "\n".join(p.text for p in pages)
    # doc_id keyed on name + content so an edited file re-ingests as a new version
    # while an unchanged re-upload is a no-op upsert.
    doc_id = stable_id(path.name, content_hash(body))

    return SourceDocument(
        doc_id=doc_id,
        source=path.name,
        pages=pages,
        metadata={
            "doc_type": doc_type or _guess_doc_type(path.name),
            "file_path": str(path),
            "file_suffix": suffix,
            "page_count": len(pages),
            "content_hash": content_hash(body),
        },
    )


def load_directory(directory: str | Path, doc_type: str | None = None) -> list[SourceDocument]:
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"Not a directory: {directory}")

    docs: list[SourceDocument] = []
    for path in sorted(directory.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
            try:
                docs.append(load_document(path, doc_type=doc_type))
            except (FileNotFoundError, ValueError, OSError) as exc:
                logger.warning("Skipping %s: %s", path.name, exc)
    return docs


def _guess_doc_type(filename: str) -> str:
    name = filename.lower()
    rules = (
        (("financial", "earnings", "10k", "10-k", "balance"), "financial"),
        (("hr", "handbook", "employee", "leave", "benefit"), "hr"),
        (("aws", "cloud", "architecture", "infra", "runbook"), "engineering"),
        (("security", "policy", "compliance", "soc2", "gdpr"), "compliance"),
        (("api", "reference", "sdk", "developer"), "api"),
        (("contract", "vendor", "msa", "sla"), "legal"),
        (("incident", "postmortem", "outage"), "operations"),
        (("risk", "portfolio", "investment", "mandate"), "investment"),
    )
    for keywords, label in rules:
        if any(k in name for k in keywords):
            return label
    return "general"
